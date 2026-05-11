from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, status
from fastapi.security import OAuth2PasswordBearer
from typing import List, Optional
from pydantic import BaseModel
from jose import JWTError, jwt
from passlib.context import CryptContext
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from duckduckgo_search import DDGS
from PIL import Image
import sqlite3
import docx
import os
import easyocr
import numpy as np
import fitz
import hashlib


print("🚀 Orion Assistant backend running with Groq + Login + Fast Notes + Database...")

load_dotenv()

# =========================
# GROQ SETTINGS
# =========================

client = OpenAI(
    api_key=os.getenv("GROQ_API_KEY"),
    base_url="https://api.groq.com/openai/v1"
)

MODEL_NAME = "llama-3.1-8b-instant"


# =========================
# FASTAPI APP
# =========================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:5173",
        "http://127.0.0.1:5173",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================
# AUTH SETTINGS
# =========================

SECRET_KEY = os.getenv("SECRET_KEY", "change-this-secret-key")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

DATABASE_NAME = "ai_study_app.db"

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")


# =========================
# REQUEST MODELS
# =========================

class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str


class LoginRequest(BaseModel):
    email: str
    password: str


class SaveNoteRequest(BaseModel):
    title: str
    type: str
    content: str


# =========================
# DATABASE SETUP
# =========================

def init_db():
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS saved_notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            type TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS generation_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            type TEXT NOT NULL,
            file_names TEXT,
            page_range TEXT,
            content TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    conn.commit()
    conn.close()


init_db()


# =========================
# AUTH DB FUNCTIONS
# =========================

def get_user_by_email(email: str):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        "SELECT id, username, email, password FROM users WHERE email = ?",
        (email,)
    )

    user = cursor.fetchone()
    conn.close()

    if user:
        return {
            "id": user[0],
            "username": user[1],
            "email": user[2],
            "password": user[3]
        }

    return None


def get_user_by_id(user_id: int):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        "SELECT id, username, email FROM users WHERE id = ?",
        (user_id,)
    )

    user = cursor.fetchone()
    conn.close()

    if user:
        return {
            "id": user[0],
            "username": user[1],
            "email": user[2]
        }

    return None


def create_user(username: str, email: str, password: str):
    existing_user = get_user_by_email(email)

    if existing_user:
        raise HTTPException(
            status_code=400,
            detail="Email already registered"
        )

    password_bytes = password.encode("utf-8")

    if len(password_bytes) > 72:
        raise HTTPException(
            status_code=400,
            detail="Password must be 72 bytes or less. Please use a shorter password."
        )

    hashed_password = pwd_context.hash(password)

    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        "INSERT INTO users (username, email, password) VALUES (?, ?, ?)",
        (username, email, hashed_password)
    )

    conn.commit()
    user_id = cursor.lastrowid
    conn.close()

    return {
        "id": user_id,
        "username": username,
        "email": email
    }


def verify_password(plain_password: str, hashed_password: str):
    return pwd_context.verify(plain_password, hashed_password)


def create_access_token(data: dict):
    to_encode = data.copy()

    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)

    to_encode.update({"exp": expire})

    encoded_jwt = jwt.encode(
        to_encode,
        SECRET_KEY,
        algorithm=ALGORITHM
    )

    return encoded_jwt


def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Invalid or expired token",
        headers={"WWW-Authenticate": "Bearer"},
    )

    try:
        payload = jwt.decode(
            token,
            SECRET_KEY,
            algorithms=[ALGORITHM]
        )

        user_id = payload.get("sub")

        if user_id is None:
            raise credentials_exception

        user_id = int(user_id)

    except JWTError:
        raise credentials_exception

    user = get_user_by_id(user_id)

    if user is None:
        raise credentials_exception

    return user


# =========================
# SAVED NOTES + HISTORY DB FUNCTIONS
# =========================

def create_saved_note(user_id: int, title: str, note_type: str, content: str):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO saved_notes (user_id, title, type, content)
        VALUES (?, ?, ?, ?)
        """,
        (user_id, title, note_type, content)
    )

    conn.commit()
    note_id = cursor.lastrowid
    conn.close()

    return note_id


def get_saved_notes_by_user(user_id: int):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT id, title, type, content, created_at
        FROM saved_notes
        WHERE user_id = ?
        ORDER BY id DESC
        """,
        (user_id,)
    )

    rows = cursor.fetchall()
    conn.close()

    notes = []

    for row in rows:
        notes.append({
            "id": row[0],
            "title": row[1],
            "type": row[2],
            "content": row[3],
            "date": row[4],
        })

    return notes


def delete_saved_note_by_id(user_id: int, note_id: int):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        DELETE FROM saved_notes
        WHERE id = ? AND user_id = ?
        """,
        (note_id, user_id)
    )

    deleted = cursor.rowcount

    conn.commit()
    conn.close()

    return deleted > 0


def create_generation_history(
    user_id: int,
    generation_type: str,
    file_names: str,
    page_range: str,
    content: str
):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO generation_history (user_id, type, file_names, page_range, content)
        VALUES (?, ?, ?, ?, ?)
        """,
        (user_id, generation_type, file_names, page_range, content)
    )

    conn.commit()
    history_id = cursor.lastrowid
    conn.close()

    return history_id


def get_generation_history_by_user(user_id: int):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT id, type, file_names, page_range, content, created_at
        FROM generation_history
        WHERE user_id = ?
        ORDER BY id DESC
        """,
        (user_id,)
    )

    rows = cursor.fetchall()
    conn.close()

    history = []

    for row in rows:
        history.append({
            "id": row[0],
            "type": row[1],
            "file_names": row[2],
            "page_range": row[3],
            "content": row[4],
            "date": row[5],
        })

    return history


def delete_generation_history_by_id(user_id: int, history_id: int):
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()

    cursor.execute(
        """
        DELETE FROM generation_history
        WHERE id = ? AND user_id = ?
        """,
        (history_id, user_id)
    )

    deleted = cursor.rowcount

    conn.commit()
    conn.close()

    return deleted > 0


# =========================
# OCR SETTINGS
# =========================

ocr_reader = None

OCR_CACHE_DIR = Path("ocr_cache")
OCR_CACHE_DIR.mkdir(exist_ok=True)

MAX_OCR_PDF_PAGES = 50

PDF_OCR_RENDER_SCALE = 1.0
PDF_RENDER_WORKERS = 4

OCR_BATCH_SIZE = 2
OCR_CANVAS_SIZE = 1280
OCR_MAG_RATIO = 0.7

FAST_NOTES_MAX_PAGES = 30
FAST_NOTES_CHUNK_SIZE = 3000
FAST_NOTES_OVERLAP = 200


def get_ocr_reader():
    global ocr_reader

    if ocr_reader is None:
        print("🔍 Loading OCR model... First OCR use may take time.")

        try:
            import torch
            use_gpu = torch.cuda.is_available()
        except Exception:
            use_gpu = False

        if use_gpu:
            print("✅ CUDA GPU detected. EasyOCR will use GPU.")
        else:
            print("⚠ No CUDA GPU detected. EasyOCR will use CPU.")

        ocr_reader = easyocr.Reader(["en"], gpu=use_gpu)

    return ocr_reader


def get_file_hash(content: bytes):
    return hashlib.md5(content).hexdigest()


def get_cached_ocr_text(file_hash: str):
    cache_file = OCR_CACHE_DIR / f"{file_hash}.txt"

    if cache_file.exists():
        print("⚡ Using cached OCR text")
        return cache_file.read_text(encoding="utf-8")

    return None


def save_cached_ocr_text(file_hash: str, text: str):
    cache_file = OCR_CACHE_DIR / f"{file_hash}.txt"
    cache_file.write_text(text, encoding="utf-8")


def ocr_single_image(image: Image.Image) -> str:
    image = image.convert("RGB")
    image_np = np.array(image)

    reader = get_ocr_reader()

    results = reader.readtext(
        image_np,
        detail=0,
        paragraph=False,
        batch_size=OCR_BATCH_SIZE,
        decoder="greedy",
        canvas_size=OCR_CANVAS_SIZE,
        mag_ratio=OCR_MAG_RATIO,
        text_threshold=0.5,
        low_text=0.3,
        link_threshold=0.3,
        width_ths=0.7,
        workers=0
    )

    return "\n".join(results).strip()


def ocr_pdf_pages_fast(images_with_page_numbers):
    reader = get_ocr_reader()
    final_text = ""

    for page_number, image_np in images_with_page_numbers:
        print(f"⚡ OCR running on page {page_number + 1}")

        results = reader.readtext(
            image_np,
            detail=0,
            paragraph=False,
            batch_size=OCR_BATCH_SIZE,
            decoder="greedy",
            canvas_size=OCR_CANVAS_SIZE,
            mag_ratio=OCR_MAG_RATIO,
            text_threshold=0.5,
            low_text=0.3,
            link_threshold=0.3,
            width_ths=0.7,
            workers=0
        )

        page_text = "\n".join(results).strip()

        if page_text:
            final_text += f"\n\n--- OCR Text from PDF Page {page_number + 1} ---\n"
            final_text += page_text

        print(f"✅ OCR completed for page {page_number + 1}")

    return final_text.strip()


def extract_pdf_text_fast(content: bytes, file_hash: str):
    cached_text = get_cached_ocr_text(file_hash)

    if cached_text:
        return cached_text

    pdf_document = fitz.open(stream=content, filetype="pdf")

    total_pages = len(pdf_document)
    pages_limit = min(total_pages, MAX_OCR_PDF_PAGES)

    print(f"📄 PDF pages found: {total_pages}")
    print(f"⚡ Processing first {pages_limit} pages")

    final_pages = [""] * pages_limit
    pages_needing_ocr = []

    for page_number in range(pages_limit):
        page = pdf_document[page_number]

        try:
            page_text = page.get_text("text").strip()
        except Exception:
            page_text = ""

        if len(page_text) > 50:
            final_pages[page_number] = f"\n\n--- Page {page_number + 1} ---\n{page_text}"
            print(f"✅ Selectable text found on page {page_number + 1}")
        else:
            pages_needing_ocr.append(page_number)
            print(f"⚠ Page {page_number + 1} needs OCR")

    def render_page(page_number: int):
        page = pdf_document[page_number]

        pix = page.get_pixmap(
            matrix=fitz.Matrix(PDF_OCR_RENDER_SCALE, PDF_OCR_RENDER_SCALE),
            alpha=False
        )

        image = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
        image_np = np.array(image)

        return page_number, image_np

    images_with_page_numbers = []

    if pages_needing_ocr:
        print("⚡ Rendering OCR pages in parallel...")

        with ThreadPoolExecutor(max_workers=PDF_RENDER_WORKERS) as executor:
            futures = [
                executor.submit(render_page, page_number)
                for page_number in pages_needing_ocr
            ]

            for future in as_completed(futures):
                page_number, image_np = future.result()
                images_with_page_numbers.append((page_number, image_np))

        images_with_page_numbers.sort(key=lambda x: x[0])

    pdf_document.close()

    if images_with_page_numbers:
        ocr_text = ocr_pdf_pages_fast(images_with_page_numbers)

        sections = ocr_text.split("\n\n--- OCR Text from PDF Page ")

        for section in sections:
            if not section.strip():
                continue

            try:
                page_part, text_part = section.split(" ---\n", 1)
                page_number = int(page_part.strip()) - 1

                if 0 <= page_number < len(final_pages):
                    final_pages[page_number] = f"\n\n--- Page {page_number + 1} ---\n{text_part.strip()}"

            except Exception:
                pass

    text = "\n\n".join([page for page in final_pages if page.strip()])

    if total_pages > MAX_OCR_PDF_PAGES:
        text += (
            f"\n\n--- Note: Only first {MAX_OCR_PDF_PAGES} pages were processed "
            f"out of {total_pages} pages. Use Fast Notes page range for large PDFs. ---"
        )

    print("📄 Final PDF extracted characters:", len(text))

    save_cached_ocr_text(file_hash, text.strip())

    return text.strip()


def get_range_cache_key(file_hash: str, start_page: int, end_page: int):
    return f"{file_hash}_pages_{start_page}_to_{end_page}"


def extract_pdf_text_by_page_range(
    content: bytes,
    file_hash: str,
    start_page: int = 1,
    end_page: int = 30
):
    start_page = max(1, start_page)
    end_page = max(start_page, end_page)

    if end_page - start_page + 1 > FAST_NOTES_MAX_PAGES:
        end_page = start_page + FAST_NOTES_MAX_PAGES - 1

    range_cache_key = get_range_cache_key(file_hash, start_page, end_page)

    cached_text = get_cached_ocr_text(range_cache_key)

    if cached_text:
        return cached_text

    pdf_document = fitz.open(stream=content, filetype="pdf")
    total_pages = len(pdf_document)

    start_index = start_page - 1
    end_index = min(end_page, total_pages)

    if start_index >= total_pages:
        pdf_document.close()
        return ""

    print(f"📄 PDF pages found: {total_pages}")
    print(f"⚡ Fast mode processing pages {start_page} to {end_index}")

    final_pages = [""] * (end_index - start_index)
    pages_needing_ocr = []

    for real_page_index in range(start_index, end_index):
        local_index = real_page_index - start_index
        page = pdf_document[real_page_index]

        try:
            page_text = page.get_text("text").strip()
        except Exception:
            page_text = ""

        if len(page_text) > 50:
            final_pages[local_index] = (
                f"\n\n--- Page {real_page_index + 1} ---\n{page_text}"
            )
            print(f"✅ Selectable text found on page {real_page_index + 1}")
        else:
            pages_needing_ocr.append(real_page_index)
            print(f"⚠ Page {real_page_index + 1} needs OCR")

    def render_page(page_number: int):
        page = pdf_document[page_number]

        pix = page.get_pixmap(
            matrix=fitz.Matrix(PDF_OCR_RENDER_SCALE, PDF_OCR_RENDER_SCALE),
            alpha=False
        )

        image = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
        image_np = np.array(image)

        return page_number, image_np

    images_with_page_numbers = []

    if pages_needing_ocr:
        print("⚡ Rendering selected OCR pages in parallel...")

        with ThreadPoolExecutor(max_workers=PDF_RENDER_WORKERS) as executor:
            futures = [
                executor.submit(render_page, page_number)
                for page_number in pages_needing_ocr
            ]

            for future in as_completed(futures):
                page_number, image_np = future.result()
                images_with_page_numbers.append((page_number, image_np))

        images_with_page_numbers.sort(key=lambda x: x[0])

    pdf_document.close()

    if images_with_page_numbers:
        reader = get_ocr_reader()

        for page_number, image_np in images_with_page_numbers:
            print(f"⚡ OCR running on selected page {page_number + 1}")

            results = reader.readtext(
                image_np,
                detail=0,
                paragraph=False,
                batch_size=OCR_BATCH_SIZE,
                decoder="greedy",
                canvas_size=OCR_CANVAS_SIZE,
                mag_ratio=OCR_MAG_RATIO,
                text_threshold=0.5,
                low_text=0.3,
                link_threshold=0.3,
                width_ths=0.7,
                workers=0
            )

            page_text = "\n".join(results).strip()
            local_index = page_number - start_index

            if page_text and 0 <= local_index < len(final_pages):
                final_pages[local_index] = (
                    f"\n\n--- OCR Text from Page {page_number + 1} ---\n{page_text}"
                )

            print(f"✅ OCR completed for selected page {page_number + 1}")

    text = "\n\n".join([page for page in final_pages if page.strip()])

    print("📄 Fast range extracted characters:", len(text))

    save_cached_ocr_text(range_cache_key, text.strip())

    return text.strip()


# =========================
# FILE EXTRACTION
# =========================

async def extract_text(file: UploadFile):
    print("📂 Uploaded file:", file.filename)

    if not file.filename:
        return None

    filename = file.filename.lower()
    content = await file.read()
    file_hash = get_file_hash(content)
    file_bytes = BytesIO(content)

    try:
        if filename.endswith(".pdf"):
            text = extract_pdf_text_fast(content, file_hash)

            if text:
                return text.strip()

            return None

        elif filename.endswith(".docx"):
            document = docx.Document(file_bytes)
            text = ""

            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"

            print("DOCX extracted characters:", len(text))
            return text.strip()

        elif filename.endswith(".txt"):
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")

            print("TXT extracted characters:", len(text))
            return text.strip()

        elif filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
            cached_text = get_cached_ocr_text(file_hash)

            if cached_text:
                return cached_text

            image = Image.open(file_bytes).convert("RGB")
            text = ocr_single_image(image)

            print("IMAGE OCR extracted characters:", len(text))

            save_cached_ocr_text(file_hash, text.strip())

            return text.strip()

        return None

    except Exception as e:
        print("❌ FILE ERROR:", e)
        return None


async def extract_text_from_files(files: List[UploadFile]):
    combined_text = ""

    for file in files:
        text = await extract_text(file)

        if text:
            combined_text += f"\n\n--- Content from {file.filename} ---\n"
            combined_text += text
        else:
            combined_text += f"\n\n--- Could not extract text from {file.filename} ---\n"

    return combined_text.strip()


async def extract_text_from_files_fast_range(
    files: List[UploadFile],
    start_page: int = 1,
    end_page: int = 30
):
    combined_text = ""

    for file in files:
        print("📂 Uploaded file:", file.filename)

        if not file.filename:
            continue

        filename = file.filename.lower()
        content = await file.read()
        file_hash = get_file_hash(content)

        if filename.endswith(".pdf"):
            text = extract_pdf_text_by_page_range(
                content=content,
                file_hash=file_hash,
                start_page=start_page,
                end_page=end_page
            )
        else:
            text = ""

            if filename.endswith(".docx"):
                document = docx.Document(BytesIO(content))

                for paragraph in document.paragraphs:
                    text += paragraph.text + "\n"

                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"

            elif filename.endswith(".txt"):
                try:
                    text = content.decode("utf-8")
                except UnicodeDecodeError:
                    text = content.decode("latin-1")

            elif filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
                image = Image.open(BytesIO(content)).convert("RGB")
                text = ocr_single_image(image)

        if text:
            combined_text += f"\n\n--- Content from {file.filename} ---\n{text}"
        else:
            combined_text += f"\n\n--- Could not extract text from {file.filename} ---\n"

    return combined_text.strip()


# =========================
# AI FUNCTIONS
# =========================

def ask_ai(prompt):
    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert academic study assistant."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7
        )

        return response.choices[0].message.content

    except Exception as e:
        print("❌ GROQ ERROR:", e)
        return f"Groq Error: {str(e)}"


def search_internet(query, max_results=5):
    try:
        results_text = ""

        with DDGS() as ddgs:
            results = ddgs.text(query, max_results=max_results)

            for i, result in enumerate(results, start=1):
                title = result.get("title", "No title")
                body = result.get("body", "No description")
                href = result.get("href", "")

                results_text += f"""
Source {i}:
Title: {title}
Summary: {body}
Link: {href}
"""

        if not results_text.strip():
            return "No useful internet search results found."

        return results_text.strip()

    except Exception as e:
        print("❌ INTERNET SEARCH ERROR:", e)
        return "Internet search failed or is unavailable."


# =========================
# LARGE TEXT CHUNKING
# =========================

def split_text_into_chunks(text: str, chunk_size: int = 3500, overlap: int = 300):
    text = text.strip()

    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start = end - overlap

        if start >= len(text):
            break

    return chunks


def split_fast_notes_chunks(text: str):
    text = text.strip()

    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = start + FAST_NOTES_CHUNK_SIZE
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start = end - FAST_NOTES_OVERLAP

        if start >= len(text):
            break

    return chunks


def combine_notes_in_small_batches(chunk_notes, batch_size: int = 3):
    if not chunk_notes:
        return "No notes were generated."

    if len(chunk_notes) == 1:
        return chunk_notes[0]

    combined_batches = []
    total_batches = (len(chunk_notes) + batch_size - 1) // batch_size

    for i in range(0, len(chunk_notes), batch_size):
        batch = chunk_notes[i:i + batch_size]

        print(f"🧩 Combining notes batch {i // batch_size + 1}/{total_batches}")

        batch_text = "\n\n".join(batch)

        combined = ask_ai(
            f"""
You are an expert academic study assistant.

Combine the following notes into one clean section.

Rules:
- Remove repeated points.
- Keep all important concepts.
- Keep it exam-focused.
- Use clear headings if useful.
- Do not add unrelated information.
- Do not make it too short.
- Ignore unclear OCR text.
- Use paragraphs plus point-wise clarity.

Notes to combine:
{batch_text[:9000]}
"""
        )

        combined_batches.append(combined)

    if len(combined_batches) == 1:
        return combined_batches[0]

    print("🧩 Creating final combined notes safely...")

    final_input = "\n\n".join(combined_batches)

    final_notes = ask_ai(
        f"""
You are an expert academic study assistant.

Create final study notes from the combined sections below.

Rules:
- Keep the notes detailed but not repetitive.
- Preserve all important exam points.
- Organize with suitable headings.
- Use paragraph explanation plus point-wise clarity.
- Do not hallucinate.
- Do not add unrelated information.
- Make it useful for B.Tech exam preparation.

At the end, include:

What you can do next:
- Give 2 to 3 useful follow-up suggestions.

Combined Sections:
{final_input[:9000]}
"""
    )

    return final_notes


def ask_ai_large_notes(text: str):
    chunks = split_text_into_chunks(
        text=text,
        chunk_size=3500,
        overlap=300
    )

    if not chunks:
        return "No readable content was found in the uploaded material."

    print(f"📚 Total chunks created for notes: {len(chunks)}")

    chunk_notes = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Generating notes for chunk {index}/{len(chunks)}")

        chunk_result = ask_ai(
            f"""
You are an expert academic study assistant.

You are reading PART {index} of {len(chunks)} from the uploaded study material.

Create accurate exam-focused notes from this part.

Rules:
- Use only readable and meaningful content.
- Do not guess from broken OCR words.
- Do not invent topics.
- Do not make unsupported claims.
- Preserve definitions, examples, classifications, and key points.
- Write in a B.Tech/student-friendly style.
- Use paragraphs plus point-wise clarity.
- Do not add "What you can do next" here.

Content Part {index}:
{chunk}
"""
        )

        chunk_notes.append(
            f"===== NOTES FROM PART {index} =====\n{chunk_result}"
        )

    final_notes = combine_notes_in_small_batches(
        chunk_notes=chunk_notes,
        batch_size=3
    )

    return final_notes


def ask_ai_fast_notes(text: str, start_page: int, end_page: int):
    chunks = split_fast_notes_chunks(text)

    if not chunks:
        return "No readable content was found in the selected pages."

    print(f"⚡ Fast notes chunks created: {len(chunks)}")

    all_notes = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Fast notes generation chunk {index}/{len(chunks)}")

        result = ask_ai(
            f"""
You are an expert academic study assistant.

Create fast, accurate, exam-focused notes from this selected textbook section.

Selected page range:
Pages {start_page} to {end_page}

Rules:
- Use only readable and meaningful content.
- Do not guess from broken OCR words.
- Do not invent topics.
- Keep the notes useful for exam preparation.
- Use headings where useful.
- Use short paragraphs plus bullet points.
- Include important definitions, classifications, examples, and key points.
- Do not make it too long.
- Do not add unrelated information.
- Do not include "What you can do next" in this chunk.

Content chunk {index} of {len(chunks)}:
{chunk}
"""
        )

        all_notes.append(f"\n\n## Notes Part {index}\n{result}")

    final_notes = "\n".join(all_notes)

    final_notes += f"""

---

## What you can do next

- Generate MCQs from pages {start_page} to {end_page}.
- Generate 3-mark, 5-mark, and 20-mark questions from this same section.
- Continue with the next page range, for example pages {end_page + 1} to {end_page + 30}.
"""

    return final_notes.strip()


def ask_ai_large_summary(text: str):
    chunks = split_text_into_chunks(
        text=text,
        chunk_size=3500,
        overlap=300
    )

    if not chunks:
        return "No readable content was found in the uploaded material."

    print(f"📚 Total chunks created for summary: {len(chunks)}")

    partial_summaries = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Generating summary for chunk {index}/{len(chunks)}")

        partial_summary = ask_ai(
            f"""
Create clear revision summary points from this part of the uploaded material.

Rules:
- Use only meaningful readable content.
- Do not guess from broken OCR text.
- Keep important definitions, concepts, examples, and exam points.
- Make it useful for quick revision.
- Do not add random information.

Part {index} of {len(chunks)}:
{chunk}
"""
        )

        partial_summaries.append(
            f"===== SUMMARY FROM PART {index} =====\n{partial_summary}"
        )

    final_summary = combine_notes_in_small_batches(
        chunk_notes=partial_summaries,
        batch_size=3
    )

    return final_summary


def generate_questions_from_full_text(text: str):
    chunks = split_text_into_chunks(text)

    if not chunks:
        return "No readable content was found in the uploaded material."

    partial_questions = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Generating questions from chunk {index}/{len(chunks)}")

        result = ask_ai(
            f"""
You are an expert university exam question paper generator.

Generate important exam questions from this part of the uploaded material.

Rules:
- Do NOT give answers.
- Do NOT give explanations.
- Do NOT guess from broken OCR words.
- Generate useful 3-mark, 5-mark, and 20-mark style questions.
- Avoid repeated questions.
- Use clear university-level language.

Part {index} of {len(chunks)}:
{chunk}
"""
        )

        partial_questions.append(result)

    combined = "\n\n".join(partial_questions)

    final_questions = ask_ai(
        f"""
You are an expert university exam question paper generator.

From the partial questions below, create the final question paper.

Create the output exactly in this format:

SECTION A: 3-Mark Questions
Generate exactly 10 questions.

SECTION B: 5-Mark Questions
Generate exactly 10 questions.

SECTION C: 20-Mark Questions
Generate exactly 10 questions.

Rules:
- Do NOT give answers.
- Do NOT give explanations.
- Avoid repeated questions.
- Cover the major topics from the uploaded material.
- Keep questions exam-oriented and university-level.

Partial Questions:
{combined[:9000]}
"""
    )

    return final_questions


def generate_mcq_from_full_text(text: str):
    chunks = split_text_into_chunks(text)

    if not chunks:
        return "No readable content was found in the uploaded material."

    partial_mcqs = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Generating MCQs from chunk {index}/{len(chunks)}")

        result = ask_ai(
            f"""
You are an exam-focused MCQ generator.

Create MCQs only from this readable part of the material.

Rules:
- Do NOT guess from broken OCR words.
- Do NOT invent topics.
- Create meaningful MCQs only.
- Each MCQ should have 4 options.
- Mark the correct answer.
- Add a short explanation only when useful.

Part {index} of {len(chunks)}:
{chunk}
"""
        )

        partial_mcqs.append(result)

    combined = "\n\n".join(partial_mcqs)

    final_mcqs = ask_ai(
        f"""
You are an exam-focused MCQ generator.

From the partial MCQs below, create a clean final MCQ set.

Rules:
- Remove repeated MCQs.
- Keep important concepts.
- Use clear numbering.
- Each MCQ must have 4 options.
- Mark the correct answer.
- Add short explanation only when useful.
- Do not include questions based on unclear OCR text.

At the end, include:
What you can do next:
- Give 2 to 3 useful suggestions for improving revision.

Partial MCQs:
{combined[:9000]}
"""
    )

    return final_mcqs


def generate_advanced_questions_from_full_text(text: str):
    chunks = split_text_into_chunks(text)

    if not chunks:
        return "No readable content was found in the uploaded material."

    partial_questions = []

    for index, chunk in enumerate(chunks, start=1):
        print(f"🧠 Generating advanced questions from chunk {index}/{len(chunks)}")

        result = ask_ai(
            f"""
You are an expert university-level advanced question generator.

Generate advanced questions from this part of the uploaded material.

Rules:
- Do NOT give full answers.
- Do NOT generate MCQs.
- Do NOT guess from broken OCR words.
- Include important advanced questions, related out-of-syllabus questions,
  higher-order thinking questions, and small hints.
- Keep it connected to the material.

Part {index} of {len(chunks)}:
{chunk}
"""
        )

        partial_questions.append(result)

    combined = "\n\n".join(partial_questions)

    final_advanced = ask_ai(
        f"""
You are an expert university-level advanced question generator.

From the partial advanced questions below, create one final clean set.

The output should contain ONLY these sections:

1. Important Advanced Questions
2. Out-of-Syllabus but Related Questions
3. Higher-Order Thinking Questions
4. Small Hints

Rules:
- Remove repeated questions.
- Do not give full answers.
- Do not include MCQs.
- Keep questions exam-focused and useful.
- Small hints should be short.

Partial Advanced Questions:
{combined[:9000]}
"""
    )

    return final_advanced


def save_history_for_generation(
    user_id: int,
    generation_type: str,
    files: List[UploadFile],
    page_range: str,
    content: str
):
    file_names = ", ".join([file.filename for file in files if file.filename])

    create_generation_history(
        user_id=user_id,
        generation_type=generation_type,
        file_names=file_names,
        page_range=page_range,
        content=content
    )


# =========================
# AUTH ROUTES
# =========================

@app.post("/register")
async def register_user(data: RegisterRequest):
    if not data.username.strip():
        raise HTTPException(
            status_code=400,
            detail="Username is required"
        )

    if not data.email.strip():
        raise HTTPException(
            status_code=400,
            detail="Email is required"
        )

    if len(data.password) < 6:
        raise HTTPException(
            status_code=400,
            detail="Password must be at least 6 characters"
        )

    if len(data.password.encode("utf-8")) > 72:
        raise HTTPException(
            status_code=400,
            detail="Password must be 72 bytes or less. Please use a shorter password."
        )

    user = create_user(
        username=data.username.strip(),
        email=data.email.strip().lower(),
        password=data.password
    )

    token = create_access_token(
        data={"sub": str(user["id"])}
    )

    return {
        "message": "Registration successful",
        "access_token": token,
        "token_type": "bearer",
        "user": user
    }


@app.post("/login")
async def login_user(data: LoginRequest):
    user = get_user_by_email(data.email.strip().lower())

    if not user:
        raise HTTPException(
            status_code=401,
            detail="Invalid email or password"
        )

    if not verify_password(data.password, user["password"]):
        raise HTTPException(
            status_code=401,
            detail="Invalid email or password"
        )

    token = create_access_token(
        data={"sub": str(user["id"])}
    )

    return {
        "message": "Login successful",
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user["id"],
            "username": user["username"],
            "email": user["email"]
        }
    }


@app.get("/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    return {
        "user": current_user
    }


# =========================
# SAVED NOTES ROUTES
# =========================

@app.post("/save-note")
async def save_note(
    data: SaveNoteRequest,
    current_user: dict = Depends(get_current_user)
):
    if not data.content.strip():
        raise HTTPException(
            status_code=400,
            detail="Cannot save empty note"
        )

    title = data.title.strip() or data.type.strip() or "Untitled Note"
    note_type = data.type.strip() or "Note"

    note_id = create_saved_note(
        user_id=current_user["id"],
        title=title,
        note_type=note_type,
        content=data.content
    )

    return {
        "message": "Note saved successfully",
        "note": {
            "id": note_id,
            "title": title,
            "type": note_type,
            "content": data.content
        }
    }


@app.get("/saved-notes")
async def get_saved_notes(
    current_user: dict = Depends(get_current_user)
):
    notes = get_saved_notes_by_user(current_user["id"])

    return {
        "saved_notes": notes
    }


@app.delete("/saved-notes/{note_id}")
async def delete_saved_note(
    note_id: int,
    current_user: dict = Depends(get_current_user)
):
    deleted = delete_saved_note_by_id(
        user_id=current_user["id"],
        note_id=note_id
    )

    if not deleted:
        raise HTTPException(
            status_code=404,
            detail="Saved note not found"
        )

    return {
        "message": "Saved note deleted successfully"
    }


# =========================
# GENERATION HISTORY ROUTES
# =========================

@app.get("/generation-history")
async def get_generation_history(
    current_user: dict = Depends(get_current_user)
):
    history = get_generation_history_by_user(current_user["id"])

    return {
        "generation_history": history
    }


@app.delete("/generation-history/{history_id}")
async def delete_generation_history(
    history_id: int,
    current_user: dict = Depends(get_current_user)
):
    deleted = delete_generation_history_by_id(
        user_id=current_user["id"],
        history_id=history_id
    )

    if not deleted:
        raise HTTPException(
            status_code=404,
            detail="History item not found"
        )

    return {
        "message": "History deleted successfully"
    }


# =========================
# MAIN ROUTES
# =========================

@app.get("/")
async def root():
    return {"message": "Orion Assistant Backend Running 🚀"}


@app.post("/generate-notes-fast")
async def notes_fast(
    files: List[UploadFile] = File(...),
    start_page: int = Form(1),
    end_page: int = Form(30),
    current_user: dict = Depends(get_current_user)
):
    if end_page < start_page:
        return {"error": "End page must be greater than or equal to start page."}

    if end_page - start_page + 1 > FAST_NOTES_MAX_PAGES:
        end_page = start_page + FAST_NOTES_MAX_PAGES - 1

    text = await extract_text_from_files_fast_range(
        files=files,
        start_page=start_page,
        end_page=end_page
    )

    if not text:
        return {"error": "Could not extract text from selected pages."}

    print(f"👤 User {current_user['email']} is generating FAST notes")
    print(f"📄 Selected pages: {start_page} to {end_page}")
    print("📄 Extracted characters:", len(text))

    result = ask_ai_fast_notes(
        text=text,
        start_page=start_page,
        end_page=end_page
    )

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="Fast Notes",
        files=files,
        page_range=f"{start_page}-{end_page}",
        content=result
    )

    return {
        "notes": result,
        "processed_pages": {
            "start_page": start_page,
            "end_page": end_page
        }
    }


@app.post("/generate-notes")
async def notes(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    text = await extract_text_from_files(files)

    if not text:
        return {"error": "Could not extract text from uploaded files"}

    print(f"👤 User {current_user['email']} is generating notes")
    print("📄 Total extracted characters for notes:", len(text))

    result = ask_ai_large_notes(text)

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="Important Notes",
        files=files,
        page_range="Full document",
        content=result
    )

    return {"notes": result}


@app.post("/generate-summary")
async def summary(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    text = await extract_text_from_files(files)

    if not text:
        return {"error": "Could not extract text from uploaded files"}

    print(f"👤 User {current_user['email']} is generating summary")
    print("📄 Total extracted characters for summary:", len(text))

    result = ask_ai_large_summary(text)

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="Summary",
        files=files,
        page_range="Full document",
        content=result
    )

    return {"summary": result}


@app.post("/generate-questions")
async def questions(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    text = await extract_text_from_files(files)

    if not text:
        return {"error": "Could not extract text from uploaded files"}

    print(f"👤 User {current_user['email']} is generating questions")
    print("📄 Total extracted characters for questions:", len(text))

    result = generate_questions_from_full_text(text)

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="Exam Questions",
        files=files,
        page_range="Full document",
        content=result
    )

    return {"questions": result}


@app.post("/generate-mcq")
async def mcq(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    text = await extract_text_from_files(files)

    if not text:
        return {"error": "Could not extract text from uploaded files"}

    print(f"👤 User {current_user['email']} is generating MCQs")
    print("📄 Total extracted characters for MCQ:", len(text))

    result = generate_mcq_from_full_text(text)

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="MCQs",
        files=files,
        page_range="Full document",
        content=result
    )

    return {"mcqs": result}


@app.post("/generate-advanced-questions")
async def advanced_questions(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    text = await extract_text_from_files(files)

    if not text:
        return {"error": "Could not extract text from uploaded files"}

    print(f"👤 User {current_user['email']} is generating advanced questions")
    print("📄 Total extracted characters for advanced questions:", len(text))

    result = generate_advanced_questions_from_full_text(text)

    save_history_for_generation(
        user_id=current_user["id"],
        generation_type="Advanced Questions",
        files=files,
        page_range="Full document",
        content=result
    )

    return {"advanced_questions": result}


@app.post("/chat")
async def chat_with_notes(
    files: Optional[List[UploadFile]] = File(None),
    question: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    notes_text = ""

    if files:
        extracted = await extract_text_from_files(files)

        if extracted:
            notes_text = extracted[:9000]

    internet_results = search_internet(question, max_results=5)

    if notes_text:
        notes_section = f"""
Uploaded Study Material:
{notes_text}
"""
        source_instruction = """
The student has uploaded study material. First use the uploaded notes to answer.
If the notes are not enough, add useful extra information from internet search results and general academic knowledge.
If OCR text looks corrupted, ignore unclear broken words and use only readable meaningful text.
"""
    else:
        notes_section = ""
        source_instruction = """
The student has not uploaded study material for this question.
Do NOT say "no notes were provided" in the answer.
Simply answer naturally using internet search results and general academic knowledge.
"""

    prompt = f"""
You are Orion AI, an intelligent academic study assistant.

Answer the student's question in a natural, helpful, and confident way.

Writing style:
- Be clear and student-friendly.
- Give a proper introduction before explaining.
- Use paragraphs when explanation is needed.
- Use bullet points only when they improve clarity.
- Be precise and accurate.
- Avoid unnecessary long answers.
- Do not hallucinate.
- Do not guess from broken OCR words.
- If the question is exam-related, make the answer exam-focused.

{source_instruction}

{notes_section}

Internet Search Results:
{internet_results}

Student Question:
{question}

Give the best possible answer.

At the end, include:

What you can do next:
- Give 1 or 2 useful follow-up questions or suggestions.
"""

    print(f"👤 User {current_user['email']} is using chat")

    answer = ask_ai(prompt)

    create_generation_history(
        user_id=current_user["id"],
        generation_type="Chat Answer",
        file_names=", ".join([file.filename for file in files if file and file.filename]) if files else "",
        page_range="Chat",
        content=answer
    )

    return {"answer": answer}